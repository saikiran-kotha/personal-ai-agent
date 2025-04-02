import fitz  # PyMuPDF
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from langchain.schema import Document

def load_pdf(file_path):
    pdf_document = fitz.open(file_path)
    text = ''
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    return Document(page_content=text, metadata={"source": file_path})

def load_docx(file_path):
    doc = DocxDocument(file_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return Document(page_content=text, metadata={"source": file_path})

def load_txt(file_path):
    with open(file_path, encoding='utf-8') as file:
        text = file.read()
    return Document(page_content=text, metadata={"source": file_path})

def load_html(file_path):
    with open(file_path, encoding='utf-8') as file:
        soup = BeautifulSoup(file, 'html.parser')
        text = soup.get_text()
    return Document(page_content=text, metadata={"source": file_path})
